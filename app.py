# app.py
# Comprehensive Room-based Build Estimator with Save/Reset and full report export (Word + Excel)
# Requirements: streamlit, pandas, numpy, plotly, python-docx, openpyxl, kaleido
# Run: streamlit run app.py

import streamlit as st
import pandas as pd
import numpy as np
import json
import os
import tempfile
from datetime import datetime
import plotly.express as px
import plotly.graph_objects as go
from docx import Document
from io import BytesIO

# ------------------------------
# Page config & filenames
# ------------------------------
st.set_page_config(page_title="Comprehensive Build Estimator", layout="wide", page_icon="🏠")
MATERIALS_FILE = "materials_db.json"
PROJECTS_FILE = "projects_db.json"

# ------------------------------
# Utilities: save plotly fig to PNG (kaleido required)
# ------------------------------
def fig_to_png_bytes(fig, width=1200, height=600, scale=1):
    """
    Returns PNG bytes for a plotly figure.
    Requires kaleido package.
    """
    try:
        img_bytes = fig.to_image(format="png", width=width, height=height, scale=scale)
        return img_bytes
    except Exception as e:
        # fallback: return None if kaleido not available
        return None

# ------------------------------
# Default materials DB & room templates (unchanged)
# ------------------------------
def default_materials_list():
    return [
        # Foundation
        {"item":"Cement (50kg bag)","unit":"bag","price":125.0,"phase":"Foundation","consumption_per_m2":0.40},
        {"item":"Sand (m³)","unit":"m3","price":150.0,"phase":"Foundation","consumption_per_m2":0.03},
        {"item":"Coarse aggregate (m³)","unit":"m3","price":300.0,"phase":"Foundation","consumption_per_m2":0.02},
        {"item":"Rebar (steel) per meter","unit":"m","price":22.0,"phase":"Foundation","consumption_per_m2":1.0},
        {"item":"Binding wire (kg)","unit":"kg","price":18.0,"phase":"Foundation","consumption_per_m2":0.08},
        {"item":"Formwork plywood (m²)","unit":"m2","price":60.0,"phase":"Foundation","consumption_per_m2":0.05},
        {"item":"Formwork oil (ltr)","unit":"ltr","price":25.0,"phase":"Foundation","consumption_per_m2":0.005},

        # Structure / Masonry
        {"item":"Concrete block (unit)","unit":"unit","price":6.5,"phase":"Structure","consumption_per_m2":10.0},
        {"item":"Mortar (cement bag equiv)","unit":"bag","price":125.0,"phase":"Structure","consumption_per_m2":0.12},
        {"item":"Tie wire (kg)","unit":"kg","price":18.0,"phase":"Structure","consumption_per_m2":0.02},
        {"item":"Scaffolding (rental per m2-eq)","unit":"m2-eq","price":40.0,"phase":"Structure","consumption_per_m2":0.05},
        {"item":"Concrete pump (rental per m2-eq)","unit":"m2-eq","price":12.0,"phase":"Structure","consumption_per_m2":0.02},

        # Roofing
        {"item":"Roofing sheet (m²)","unit":"m2","price":90.0,"phase":"Roofing","consumption_per_m2":1.05},
        {"item":"Timber (lm)","unit":"lm","price":40.0,"phase":"Roofing","consumption_per_m2":0.20},
        {"item":"Roof nails (kg)","unit":"kg","price":40.0,"phase":"Roofing","consumption_per_m2":0.02},
        {"item":"Anti-rust paint (ltr)","unit":"ltr","price":70.0,"phase":"Roofing","consumption_per_m2":0.02},
        {"item":"Roofing felt (m²)","unit":"m2","price":20.0,"phase":"Roofing","consumption_per_m2":0.95},
        {"item":"Truss (lm)","unit":"lm","price":85.0,"phase":"Roofing","consumption_per_m2":0.06},
        {"item":"Gutter (m)","unit":"m","price":45.0,"phase":"Roofing","consumption_per_m2":0.05},

        # Finishing
        {"item":"Floor tiles (m²)","unit":"m2","price":80.0,"phase":"Finishing","consumption_per_m2":1.0},
        {"item":"Tile adhesive (bag)","unit":"bag","price":120.0,"phase":"Finishing","consumption_per_m2":0.03},
        {"item":"Grout (kg)","unit":"kg","price":25.0,"phase":"Finishing","consumption_per_m2":0.6},
        {"item":"Wall plaster (m²)","unit":"m2","price":25.0,"phase":"Finishing","consumption_per_m2":1.0},
        {"item":"POP (kg)","unit":"kg","price":8.0,"phase":"Finishing","consumption_per_m2":0.8},
        {"item":"Paint (ltr)","unit":"ltr","price":40.0,"phase":"Finishing","consumption_per_m2":0.08},
        {"item":"Skirting (lm)","unit":"lm","price":20.0,"phase":"Finishing","consumption_per_m2":0.15},
        {"item":"Ceiling board (m²)","unit":"m2","price":120.0,"phase":"Finishing","consumption_per_m2":1.0},

        # Doors & windows
        {"item":"External door (each)","unit":"each","price":1800.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Internal door (each)","unit":"each","price":700.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Window (each)","unit":"each","price":450.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Glass (m²)","unit":"m2","price":120.0,"phase":"Finishing","consumption_per_m2":0},

        # Electrical
        {"item":"Wiring roll (100m)","unit":"roll","price":450.0,"phase":"Finishing","consumption_per_m2":0.02},
        {"item":"Switch/socket (each)","unit":"each","price":40.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"DB board (each)","unit":"each","price":1200.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Light fixture (each)","unit":"each","price":150.0,"phase":"Finishing","consumption_per_m2":0},

        # Plumbing
        {"item":"PVC pipe (10ft)","unit":"piece","price":150.0,"phase":"Finishing","consumption_per_m2":0.02},
        {"item":"Toilet (WC) (each)","unit":"each","price":700.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Wash basin (each)","unit":"each","price":450.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Shower set (each)","unit":"each","price":600.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Kitchen sink (each)","unit":"each","price":600.0,"phase":"Finishing","consumption_per_m2":0},
        {"item":"Kitchen cabinet (per lm)","unit":"lm","price":900.0,"phase":"Finishing","consumption_per_m2":0},

        # External / misc
        {"item":"Fence block (unit)","unit":"unit","price":6.0,"phase":"External","consumption_per_m2":0},
        {"item":"Gate (each)","unit":"each","price":4000.0,"phase":"External","consumption_per_m2":0},
        {"item":"Paving stone (m²)","unit":"m2","price":120.0,"phase":"External","consumption_per_m2":0},
        {"item":"Water tank (5000L)","unit":"each","price":5000.0,"phase":"External","consumption_per_m2":0},

        # Fasteners
        {"item":"Nails / screws (kg)","unit":"kg","price":30.0,"phase":"Misc","consumption_per_m2":0.04},
        {"item":"Adhesive / sealant (ltr)","unit":"ltr","price":60.0,"phase":"Misc","consumption_per_m2":0.02},
        {"item":"Paint brushes / rollers (set)","unit":"set","price":120.0,"phase":"Misc","consumption_per_m2":0.001},

        # Labour (GHS per m2)
        {"item":"Masonry labour (GHS/m2)","unit":"GHS/m2","price":120.0,"phase":"Labor","consumption_per_m2":1.0},
        {"item":"Carpentry labour (GHS/m2)","unit":"GHS/m2","price":80.0,"phase":"Labor","consumption_per_m2":1.0},
        {"item":"Electrical labour (GHS/m2)","unit":"GHS/m2","price":40.0,"phase":"Labor","consumption_per_m2":1.0},
        {"item":"Plumbing labour (GHS/m2)","unit":"GHS/m2","price":35.0,"phase":"Labor","consumption_per_m2":1.0},
    ]

DEFAULT_ROOM_TEMPLATES = {
    "Master Bedroom": {"area_m2":18, "fixtures": {"Internal door":1, "Window (each)":2, "Light fixture":2}},
    "Standard Bedroom": {"area_m2":12, "fixtures": {"Internal door":1, "Window (each)":1, "Light fixture":2}},
    "Bathroom (full)": {"area_m2":6, "fixtures": {"Toilet (WC) (each)":1, "Wash basin (each)":1, "Shower set (each)":1, "Light fixture":1}},
    "Bathroom (half)": {"area_m2":3, "fixtures": {"Toilet (WC) (each)":1, "Wash basin (each)":1, "Light fixture":1}},
    "Kitchen (main)": {"area_m2":10, "fixtures": {"Kitchen sink (each)":1, "Kitchen cabinet (per lm)":3, "Light fixture":2}},
    "Living Room": {"area_m2":20, "fixtures": {"Internal door":1, "Window (each)":3, "Light fixture":3}},
    "Dining Room": {"area_m2":12, "fixtures": {"Light fixture":2, "Window (each)":2}},
    "Garage (1 car)": {"area_m2":18, "fixtures": {"Gate (each)":1}},
    "Verandah": {"area_m2":8, "fixtures": {}},
    "Store / Pantry": {"area_m2":6, "fixtures": {}}
}

FIXTURE_ITEM_MAP = {
    "Internal door": "Internal door (each)",
    "Window (each)": "Window (each)",
    "Light fixture": "Light fixture (each)",
    "Toilet (WC) (each)": "Toilet (WC) (each)",
    "Wash basin (each)": "Wash basin (each)",
    "Shower set (each)": "Shower set (each)",
    "Kitchen sink (each)": "Kitchen sink (each)",
    "Kitchen cabinet (per lm)": "Kitchen cabinet (per lm)",
    "Gate (each)": "Gate (each)"
}

# ------------------------------
# DB helpers
# ------------------------------
def load_materials_db():
    if os.path.exists(MATERIALS_FILE):
        try:
            with open(MATERIALS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            df = pd.DataFrame(data)
            for col in ["item","unit","price","phase","consumption_per_m2"]:
                if col not in df.columns:
                    df[col] = 0
            return df
        except Exception:
            return pd.DataFrame(default_materials_list())
    else:
        return pd.DataFrame(default_materials_list())

def save_materials_db(df):
    recs = df.to_dict(orient="records")
    with open(MATERIALS_FILE, "w", encoding="utf-8") as f:
        json.dump(recs, f, indent=2)
    st.success(f"Saved materials DB to {MATERIALS_FILE}")

def load_projects_db():
    if os.path.exists(PROJECTS_FILE):
        try:
            with open(PROJECTS_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            return data
        except Exception:
            return {}
    else:
        return {}

def save_projects_db(obj):
    with open(PROJECTS_FILE, "w", encoding="utf-8") as f:
        json.dump(obj, f, indent=2)
    st.success(f"Saved projects to {PROJECTS_FILE}")

# ------------------------------
# BOQ builder (unchanged)
# ------------------------------
def build_boq(counts, templates, materials_df, quality_multiplier=1.0, location_multiplier=1.0):
    mat = materials_df.copy()
    mat["price"] = pd.to_numeric(mat["price"], errors="coerce").fillna(0.0) * location_multiplier
    mat["consumption_per_m2"] = pd.to_numeric(mat.get("consumption_per_m2",0), errors="coerce").fillna(0.0) * quality_multiplier

    total_area = 0.0
    for room, cnt in counts.items():
        area = templates.get(room, {}).get("area_m2", 0.0)
        total_area += cnt * area

    qty_map = {}
    for _, row in mat.iterrows():
        qty_map[row["item"]] = row["consumption_per_m2"] * total_area

    item_lookup = {s.lower(): s for s in mat["item"].tolist()}

    def find_match(key):
        mapped = FIXTURE_ITEM_MAP.get(key)
        if mapped and mapped.lower() in item_lookup:
            return item_lookup[mapped.lower()]
        if key.lower() in item_lookup:
            return item_lookup[key.lower()]
        for k in item_lookup:
            if key.lower() in k:
                return item_lookup[k]
        return None

    for room, cnt in counts.items():
        if cnt <= 0: continue
        tpl = templates.get(room, {})
        fixtures = tpl.get("fixtures", {})
        for fk, fq in fixtures.items():
            if fq <= 0: continue
            matched = find_match(fk)
            if matched:
                qty_map[matched] = qty_map.get(matched, 0.0) + fq * cnt
            else:
                qty_map[f"{fk} (user)"] = qty_map.get(f"{fk} (user)", 0.0) + fq * cnt

    rows = []
    for item_name, qty in qty_map.items():
        if item_name in mat["item"].values:
            r = mat[mat["item"] == item_name].iloc[0]
            unit = r["unit"]
            unit_price = float(r["price"])
            phase = r.get("phase","Misc")
        else:
            unit = "each"
            unit_price = 0.0
            phase = "Misc"
        total_cost = qty * unit_price
        rows.append({"item":item_name, "unit":unit, "total_qty":qty, "unit_price":unit_price, "total_cost":total_cost, "phase":phase})

    boq_df = pd.DataFrame(rows)
    if boq_df.empty:
        return boq_df
    boq_df = boq_df.groupby(["item","unit","unit_price","phase"], as_index=False).agg({"total_qty":"sum","total_cost":"sum"})
    return boq_df

# ------------------------------
# Session init
# ------------------------------
if "materials_df" not in st.session_state:
    st.session_state["materials_df"] = load_materials_db()
if "room_templates" not in st.session_state:
    st.session_state["room_templates"] = DEFAULT_ROOM_TEMPLATES.copy()
if "projects" not in st.session_state:
    st.session_state["projects"] = load_projects_db()
# navigation/run flags
for k in ["page","go_to_projection","go_to_affordability","run_projection","projection_ran"]:
    if k not in st.session_state:
        st.session_state[k] = False if k != "page" else "Inputs"

# Handle navigation flags BEFORE widgets are created
if st.session_state.get("go_to_projection", False):
    st.session_state["page"] = "Projection"
    st.session_state["go_to_projection"] = False
if st.session_state.get("go_to_affordability", False):
    st.session_state["page"] = "Affordability"
    st.session_state["go_to_affordability"] = False

# Sidebar navigator (key=page)
st.sidebar.title("Navigation")
st.sidebar.radio("Go to", ["Inputs","Materials DB","Projection","Affordability","Export & Save"], key="page")

# ------------------------------
# Inputs page
# ------------------------------
if st.sidebar.button("📊 Quick -> Projection"):
    st.session_state["run_projection"] = True
    st.session_state["go_to_projection"] = True
    st.rerun()

if st.session_state["page"] == "Inputs":
    st.header("Project Inputs — Rooms & Features")
    st.markdown("Enter counts for rooms you need. Edit default area & fixtures per room if desired.")
    templates = st.session_state["room_templates"]

    project_name = st.text_input("Project name", value=st.session_state.get("current_project_name","My Project"))

    st.subheader("Room Counts")
    cols = st.columns(3)
    counts = {}
    i = 0
    for room in templates.keys():
        key = f"cnt_{room}"
        default = 1 if room in ["Standard Bedroom"] else 0
        counts[room] = cols[i % 3].number_input(room, min_value=0, value=st.session_state.get(key, default), key=key)
        i += 1

    st.subheader("Edit room templates (area & fixtures)")
    for room, tpl in templates.items():
        with st.expander(f"{room} (area {tpl['area_m2']} m²)"):
            a_key = f"area_{room}"
            new_area = st.number_input(f"{room} area (m²)", min_value=1.0, value=float(st.session_state.get(a_key, tpl["area_m2"])), key=a_key)
            st.session_state["room_templates"][room]["area_m2"] = new_area
            st.markdown("Fixtures per one room:")
            for fk, fv in tpl["fixtures"].items():
                fk_key = f"fix_{room}_{fk}"
                new_f = st.number_input(f"{fk} (per {room})", min_value=0.0, value=float(st.session_state.get(fk_key, fv)), step=1.0, key=fk_key)
                st.session_state["room_templates"][room]["fixtures"][fk] = new_f

    st.subheader("Project quality & location multipliers")
    colq1, colq2 = st.columns(2)
    quality_multiplier = colq1.select_slider("Quality multiplier (affects consumption)", options=[0.8,1.0,1.1,1.25,1.5], value=st.session_state.get("quality_multiplier",1.0))
    location_multiplier = colq2.selectbox("Location multiplier (affects prices)", options=[0.9,1.0,1.1,1.2], index=[0.9,1.0,1.1,1.2].index(st.session_state.get("location_multiplier",1.0)))
    st.session_state["quality_multiplier"] = quality_multiplier
    st.session_state["location_multiplier"] = location_multiplier

    col_save, col_reset, col_proj = st.columns([1,1,1])
    with col_save:
        if st.button("💾 Save Inputs / Create Project"):
            project = {
                "name": project_name,
                "created_at": datetime.utcnow().isoformat(),
                "counts": counts,
                "room_templates": st.session_state["room_templates"],
                "quality_multiplier": quality_multiplier,
                "location_multiplier": location_multiplier
            }
            st.session_state["projects"][project_name] = project
            save_projects_db(st.session_state["projects"])
            st.success(f"Project '{project_name}' saved.")
            st.session_state["current_project_name"] = project_name
    with col_reset:
        if st.button("♻ Reset Inputs to defaults"):
            for room in DEFAULT_ROOM_TEMPLATES:
                key = f"cnt_{room}"
                st.session_state[key] = 0
                st.session_state["room_templates"][room] = DEFAULT_ROOM_TEMPLATES[room].copy()
            st.session_state["quality_multiplier"] = 1.0
            st.session_state["location_multiplier"] = 1.0
            st.success("Inputs reset in session.")
    with col_proj:
        if st.button("🚀 Go to Projection"):
            st.session_state["unsaved_counts"] = counts
            st.session_state["go_to_projection"] = True
            st.rerun()

    st.markdown("---")
    st.subheader("Saved Projects")
    if st.session_state["projects"]:
        proj_names = list(st.session_state["projects"].keys())
        sel = st.selectbox("Load existing project", options=["(none)"] + proj_names)
        if sel and sel != "(none)":
           # if st.button("Load project into inputs"):
               # proj = st.session_state["projects"][sel]
               # st.session_state["current_project_name"] = sel
               # for room, cnt in proj["counts"].items():
               #     cnt = st.number_input(f"Count {room}", key=f"cnt_{room}", value=cnt)


# NO assignment to st.session_state["cnt_Master Bedroom"] here — the widget updates it automatically

               # st.session_state["room_templates"] = proj.get("room_templates", st.session_state["room_templates"])
                #st.session_state["quality_multiplier"] = proj.get("quality_multiplier",1.0)
               # st.session_state["location_multiplier"] = proj.get("location_multiplier",1.0)
              #  st.rerun()
            if st.button("Delete project"):
                if sel in st.session_state["projects"]:
                    del st.session_state["projects"][sel]
                    save_projects_db(st.session_state["projects"])
                    st.success(f"Deleted project {sel}")
    else:
        st.info("No saved projects yet. Save a project above.")

# ------------------------------
# Materials DB page
# ------------------------------
if st.session_state["page"] == "Materials DB":
    st.header("Materials Database (edit & save)")
    df = st.session_state["materials_df"].copy().reset_index(drop=True)
    st.markdown("Edit unit prices and consumption_per_m2. Use Import to load supplier CSV, Save to write to disk.")
    if hasattr(st, "data_editor"):
        edited = st.data_editor(df, num_rows="dynamic", use_container_width=True)
    else:
        edited = st.experimental_data_editor(df, num_rows="dynamic", use_container_width=True)
    st.session_state["materials_df"] = edited

    c1, c2, c3 = st.columns(3)
    with c1:
        if st.button("Save materials to JSON"):
            save_materials_db(st.session_state["materials_df"])
    with c2:
        if st.button("Reset materials to defaults (session)"):
            st.session_state["materials_df"] = pd.DataFrame(default_materials_list())
            st.success("Materials reset (session). Save to persist.")
    with c3:
        uploaded = st.file_uploader("Import materials CSV (cols: item,unit,price,phase,consumption_per_m2)", type=["csv"])
        if uploaded:
            try:
                csvdf = pd.read_csv(uploaded)
                st.session_state["materials_df"] = csvdf
                st.success("Imported CSV into session materials.")
            except Exception as e:
                st.error("Import failed: " + str(e))

# ------------------------------
# Projection page (BOQ, visuals, report export)
# ------------------------------
# If Quick from inputs set the flag, ensure we arrive here and run projection
if st.session_state.get("run_projection", False):
    st.session_state["go_to_projection"] = True
    st.session_state["run_projection"] = False
    st.rerun()

if st.session_state["page"] == "Projection":
    st.header("Projection — BOQ & Costs")
    materials_df = st.session_state["materials_df"].copy()
    templates = st.session_state["room_templates"]

    st.subheader("Select project or use current inputs")
    proj_choices = ["(Use current session inputs)"] + list(st.session_state["projects"].keys())
    sel_proj = st.selectbox("Project to project", proj_choices)
    if sel_proj != "(Use current session inputs)":
        project = st.session_state["projects"].get(sel_proj)
        counts = project["counts"]
        quality_multiplier = project.get("quality_multiplier",1.0)
        location_multiplier = project.get("location_multiplier",1.0)
    else:
        counts = st.session_state.get("unsaved_counts", None)
        if not counts:
            counts = {}
            for room in templates.keys():
                counts[room] = int(st.session_state.get(f"cnt_{room}", 0))
        quality_multiplier = st.session_state.get("quality_multiplier",1.0)
        location_multiplier = st.session_state.get("location_multiplier",1.0)

    col_run1, col_run2 = st.columns([1,1])
    with col_run1:
        if st.button("Run Projection (generate BOQ)"):
            st.session_state["last_projection_counts"] = counts
            st.session_state["last_projection_qmult"] = quality_multiplier
            st.session_state["last_projection_lmult"] = location_multiplier
            st.session_state["projection_ran"] = True
            st.rerun()
    with col_run2:
        if st.button("Quick run and go to Affordability"):
            st.session_state["last_projection_counts"] = counts
            st.session_state["last_projection_qmult"] = quality_multiplier
            st.session_state["last_projection_lmult"] = location_multiplier
            st.session_state["projection_ran"] = True
            st.session_state["go_to_affordability"] = True
            st.rerun()

    # Auto-run or manual run
    run_now = st.session_state.get("projection_ran", False) or st.button("Run Projection (now)")
    if run_now:
        boq = build_boq(counts, templates, materials_df, quality_multiplier, location_multiplier)
        if boq.empty:
            st.warning("No BOQ items generated. Ensure room counts > 0 and materials DB has items.")
        else:
            phase_totals = boq.groupby("phase", sort=False)["total_cost"].sum().reset_index().rename(columns={"total_cost":"phase_cost"})
            base_total = boq["total_cost"].sum()
            extra_margin = st.slider("Extra margin (%) for unknowns", 0, 30, 10)/100.0
            professional_fees = st.slider("Professional fees (%)", 0, 15, 6)/100.0
            contingency = st.slider("Contingency (%)", 0, 30, 10)/100.0

            extra_cost = base_total * extra_margin
            pro_cost = base_total * professional_fees
            cont_cost = base_total * contingency
            grand_total = base_total + extra_cost + pro_cost + cont_cost

            st.subheader("Bill of Quantities (top items)")
            st.dataframe(boq.sort_values("total_cost", ascending=False).head(400).style.format({"total_qty":"{:.2f}","unit_price":"₵ {:,.2f}","total_cost":"₵ {:,.2f}"}), use_container_width=True)

            st.subheader("Phase summary")
            phase_display = phase_totals.copy()
            phase_display["phase_cost"] = phase_display["phase_cost"].map(lambda x: f"₵ {x:,.2f}")
            st.table(phase_display.rename(columns={"phase":"Phase","phase_cost":"Cost (GHS)"}))

            c1,c2,c3,c4 = st.columns(4)
            c1.metric("Base materials cost", f"₵ {base_total:,.2f}")
            c2.metric("Extra margin", f"₵ {extra_cost:,.2f} ({extra_margin*100:.0f}%)")
            c3.metric("Professional fees", f"₵ {pro_cost:,.2f} ({professional_fees*100:.0f}%)")
            c4.metric("Contingency", f"₵ {cont_cost:,.2f} ({contingency*100:.0f}%)")

            st.markdown(f"### Grand total estimate: **₵ {grand_total:,.2f}**")

            # Visuals (plotly)
            fig_pie = px.pie(phase_totals, values="phase_cost", names="phase", title="Cost share by Phase")
            st.plotly_chart(fig_pie, use_container_width=True)
            top_items = boq.sort_values("total_cost", ascending=False).head(12)
            fig_bar = px.bar(top_items, x="item", y="total_cost", title="Top cost items", labels={"total_cost":"GHS","item":"Item"})
            st.plotly_chart(fig_bar, use_container_width=True)

            # keep last projection in session
            st.session_state["last_projection_boq"] = boq.to_dict(orient="records")
            st.session_state["last_projection_summary"] = {
                "base_total": base_total,
                "extra_margin_pct": extra_margin*100,
                "professional_fees_pct": professional_fees*100,
                "contingency_pct": contingency*100,
                "grand_total": grand_total
            }
            st.session_state["projection_ran"] = False

            # ------------------------------
            # REPORT GENERATION (WORD & EXCEL)
            # ------------------------------
            st.markdown("---")
            st.subheader("📥 Download full report (Word / Excel)")

            def create_word_report_bytes(summary, boq_df, phase_df, pie_fig, bar_fig, project_meta):
                """
                Build a Word (.docx) report containing:
                - Title & summary
                - Phase table
                - BOQ table (itemized)
                - Charts (pie & bar) embedded as images (uses plotly + kaleido)
                Returns bytes buffer
                """
                doc = Document()
                doc.add_heading(f"Build Projection Report — {project_meta.get('name','Project')}", level=1)
                doc.add_paragraph(f"Generated: {datetime.utcnow().isoformat()} (UTC)")

                doc.add_heading("Summary", level=2)
                doc.add_paragraph(f"Base total (materials): ₵ {summary['base_total']:,.2f}")
                doc.add_paragraph(f"Extra margin: {summary['extra_margin_pct']:.1f}%")
                doc.add_paragraph(f"Professional fees: {summary['professional_fees_pct']:.1f}%")
                doc.add_paragraph(f"Contingency: {summary['contingency_pct']:.1f}%")
                doc.add_paragraph(f"Grand total estimate: ₵ {summary['grand_total']:,.2f}")


                doc.add_heading("Phase totals", level=2)
                # add phase table
                pf = phase_df.copy()
                table = doc.add_table(rows=1, cols=len(pf.columns))
                hdr = table.rows[0].cells
                for idx, c in enumerate(pf.columns):
                    hdr[idx].text = str(c)
                for _, r in pf.iterrows():
                    row_cells = table.add_row().cells
                    for idx, c in enumerate(pf.columns):
                        row_cells[idx].text = str(r[c])

                doc.add_heading("Bill of Quantities (top items)", level=2)
                # limit to reasonable number of rows to include fully - include all BOQ
                bf = boq_df.copy()
                table = doc.add_table(rows=1, cols=len(bf.columns))
                hdr = table.rows[0].cells
                for idx, c in enumerate(bf.columns):
                    hdr[idx].text = str(c)
                for _, r in bf.iterrows():
                    row_cells = table.add_row().cells
                    for idx, c in enumerate(bf.columns):
                        val = r[c]
                        # format floats
                        if isinstance(val, float):
                            row_cells[idx].text = f"{val:,.2f}"
                        else:
                            row_cells[idx].text = str(val)

                doc.add_heading("Charts", level=2)
                # Save charts to temp files and embed
                tmp_files = []
                try:
                    pie_bytes = fig_to_png_bytes(pie_fig, width=900, height=600, scale=1)
                    bar_bytes = fig_to_png_bytes(bar_fig, width=1200, height=600, scale=1)
                    if pie_bytes:
                        tf = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tf.write(pie_bytes)
                        tf.flush()
                        tf.close()
                        tmp_files.append(tf.name)
                        doc.add_paragraph("Cost share by Phase:")
                        doc.add_picture(tf.name, width=None)
                    if bar_bytes:
                        tf2 = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
                        tf2.write(bar_bytes)
                        tf2.flush()
                        tf2.close()
                        tmp_files.append(tf2.name)
                        doc.add_paragraph("Top cost items:")
                        doc.add_picture(tf2.name, width=None)
                except Exception:
                    # if image embedding fails, just include note
                    doc.add_paragraph("Charts could not be embedded (kaleido missing). You can still download the Excel version with charts.")

                # cleanup temp files will be handled after saving bytes
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                # remove tmp files
                for fp in tmp_files:
                    try:
                        os.unlink(fp)
                    except Exception:
                        pass
                return buffer

            def create_excel_report_bytes(summary, boq_df, phase_df, pie_fig, bar_fig, project_meta):
                """
                Create an Excel workbook with:
                - Summary sheet
                - Phase totals sheet
                - BOQ sheet
                - Charts embedded (as images) if possible
                """
                output = BytesIO()
                with pd.ExcelWriter(output, engine="openpyxl") as writer:
                    # Summary sheet
                    s = pd.DataFrame({
                        "metric": ["base_total","extra_margin_pct","professional_fees_pct","contingency_pct","grand_total"],
                        "value": [summary["base_total"], summary["extra_margin_pct"], summary["professional_fees_pct"], summary["contingency_pct"], summary["grand_total"]]
                    })
                    s.to_excel(writer, index=False, sheet_name="Summary")
                    # Phase sheet
                    phase_df.to_excel(writer, index=False, sheet_name="PhaseTotals")
                    # BOQ sheet
                    boq_df.to_excel(writer, index=False, sheet_name="BOQ")
                    writer.close()


                output.seek(0)
                # Note: embedding images into Excel requires more work (openpyxl image insertion).
                # We'll return the workbook bytes (charts available in Word doc); users can view charts in app.
                return output

            # prepare data for report
            boq_df = boq.copy().reset_index(drop=True)
            phase_df = phase_totals.copy()
            summary = {
                "base_total": base_total,
                "extra_margin_pct": extra_margin*100,
                "professional_fees_pct": professional_fees*100,
                "contingency_pct": contingency*100,
                "grand_total": grand_total
            }
            project_meta = {"name": st.session_state.get("current_project_name", "My Project")}

            # Buttons to generate and download reports
            colw, cole = st.columns(2)
            with colw:
                word_buf = create_word_report_bytes(summary, boq_df, phase_df, fig_pie, fig_bar, project_meta)
                st.download_button(
                    label="📄 Download Full Word Report (.docx)",
                    data=word_buf,
                    file_name=f"projection_report_{project_meta['name']}_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            with cole:
                excel_buf = create_excel_report_bytes(summary, boq_df, phase_df, fig_pie, fig_bar, project_meta)
                st.download_button(
                    label="📊 Download Full Excel Report (.xlsx)",
                    data=excel_buf,
                    file_name=f"projection_report_{project_meta['name']}_{datetime.now().strftime('%Y%m%d')}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )

# If go_to_affordability was set
if st.session_state.get("go_to_affordability", False):
    st.session_state["page"] = "Affordability"
    st.session_state["go_to_affordability"] = False
    st.rerun()

# ------------------------------
# Affordability page (unchanged)
# ------------------------------
if st.session_state["page"] == "Affordability":
    st.header("Affordability — funding timeline using monthly savings")
    materials_df = st.session_state["materials_df"].copy()
    templates = st.session_state["room_templates"]

    st.subheader("Select project / use recent projection")
    if "last_projection_boq" in st.session_state:
        st.write("Using last projection BOQ from Projection page.")
        boq_df = pd.DataFrame(st.session_state["last_projection_boq"])
    else:
        st.info("Run projection first (Projection page).")
        boq_df = pd.DataFrame()

    st.write("### Funding inputs")
    one_time = st.number_input("One-time investment (₵)", min_value=0.0, value=0.0, step=100.0)
    monthly_income = st.number_input("Monthly income (₵)", min_value=0.0, value=5000.0, step=50.0)
    use_pct = st.checkbox("Save % of income", value=True)
    if use_pct:
        pct = st.slider("Percent of income to save (%)", 1, 100, 30)
        monthly_savings = monthly_income * pct / 100.0
    else:
        monthly_savings = st.number_input("Monthly savings (₵)", min_value=0.0, value=1500.0, step=50.0)
    inflation_pct = st.number_input("Annual inflation (%)", min_value=0.0, value=10.0, step=0.1)

    if not boq_df.empty:
        base_total = boq_df["total_cost"].sum()
        st.metric("Base projected cost (today)", f"₵ {base_total:,.2f}")

        phase_tot = boq_df.groupby("phase", sort=False)["total_cost"].sum().to_dict()
        phase_order = st.multiselect("Phase build order (top-to-bottom)", options=list(phase_tot.keys()), default=list(phase_tot.keys()))
        if not phase_order:
            phase_order = list(phase_tot.keys())

        remaining_upfront = one_time
        cum_month = 0.0
        results = []
        for i, ph in enumerate(phase_order, start=1):
            base_cost = phase_tot.get(ph, 0.0)
            if remaining_upfront > 0:
                applied = min(remaining_upfront, base_cost)
                base_cost_after = max(0.0, base_cost - applied)
                remaining_upfront -= applied
            else:
                base_cost_after = base_cost
            years_until_start = cum_month / 12.0
            infl_factor = (1 + inflation_pct/100.0) ** years_until_start
            inflated = base_cost_after * infl_factor
            months_needed = inflated / monthly_savings if monthly_savings > 0 else float("inf")
            start = cum_month
            end = start + months_needed
            results.append({"phase":ph, "base":base_cost, "after_upfront":base_cost_after, "inflated":inflated, "months":months_needed, "start":start, "end":end})
            cum_month = end

        res_df = pd.DataFrame(results)
        if res_df.empty:
            st.warning("No phases to show.")
        else:
            res_show = res_df.copy()
            res_show["inflated"] = res_show["inflated"].map(lambda x: f"₵ {x:,.2f}")
            res_show["months"] = res_show["months"].map(lambda x: f"{x:,.1f}" if np.isfinite(x) else "∞")
            res_show["start"] = res_show["start"].map(lambda x: f"{x:,.1f}")
            res_show["end"] = res_show["end"].map(lambda x: f"{x:,.1f}")
            st.subheader("Phase funding schedule")
            st.table(res_show[["phase","base","after_upfront","inflated","months","start","end"]].rename(columns={"base":"Base cost (GHS)","after_upfront":"After upfront (GHS)"}))

            fig = go.Figure()
            colors = px.colors.qualitative.Plotly
            for i, r in res_df.iterrows():
                length = r["end"] - r["start"]
                fig.add_trace(go.Bar(x=[length], y=[r["phase"]], base=[r["start"]], orientation='h', name=r["phase"], marker_color=colors[i % len(colors)]))
            fig.update_layout(title="Funding timeline (months)", xaxis_title="Months")
            st.plotly_chart(fig, use_container_width=True)

            total_needed = res_df["inflated"].sum()
            months_span = int(np.ceil(res_df["end"].max())) if np.isfinite(res_df["end"].max()) else 1
            months_span = max(1, months_span)
            months = np.arange(0, months_span+1)
            cumulative_savings = months * monthly_savings + one_time
            required = np.zeros_like(months, dtype=float)
            for mi, m in enumerate(months):
                req = 0.0
                for r in results:
                    if m >= r["end"]:
                        req += r["inflated"]
                    elif m > r["start"]:
                        if np.isfinite(r["months"]) and r["months"]>0:
                            frac = min(1.0, (m - r["start"]) / r["months"])
                            req += r["inflated"] * frac
                required[mi] = req
            df_cum = pd.DataFrame({"Month": months, "CumulativeSavings": cumulative_savings, "CumulativeRequired": required})
            fig2 = px.line(df_cum, x="Month", y=["CumulativeSavings","CumulativeRequired"], labels={"value":"GHS"})
            st.plotly_chart(fig2, use_container_width=True)

            final_months = res_df["end"].iloc[-1]
            if np.isfinite(final_months):
                yrs = int(final_months // 12)
                mos = int(final_months % 12)
                st.metric("Estimated completion", f"{yrs} yrs {mos} mos")
                st.metric("Total projected (inflated)", f"₵ {res_df['inflated'].sum():,.2f}")
            else:
                st.warning("Project cannot be funded with current monthly savings (infinite time). Increase savings or use financing.")

# ------------------------------
# Export & Save page (Word & Excel for DBs)
# ------------------------------
if st.session_state["page"] == "Export & Save":
    st.header("Export & Save")
    st.subheader("Materials DB")
    df_materials = st.session_state["materials_df"].copy().reset_index(drop=True)

    def export_to_word_bytes(df, title="Export"):
        doc = Document()
        doc.add_heading(title, level=1)
        table = doc.add_table(rows=1, cols=len(df.columns))
        hdr = table.rows[0].cells
        for i, c in enumerate(df.columns):
            hdr[i].text = str(c)
        for _, r in df.iterrows():
            row_cells = table.add_row().cells
            for i, c in enumerate(df.columns):
                row_cells[i].text = str(r[c])
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf

    def export_to_excel_bytes(df, sheet_name="Sheet1"):
        out = BytesIO()
        with pd.ExcelWriter(out, engine="openpyxl") as writer:
            df.to_excel(writer, index=False, sheet_name=sheet_name)
            writer.close()

        out.seek(0)
        return out

    c1, c2 = st.columns(2)
    with c1:
        word_buf = export_to_word_bytes(df_materials, "Materials Database")
        st.download_button("📄 Download Materials (Word)", word_buf, file_name=f"materials_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c2:
        xlsx_buf = export_to_excel_bytes(df_materials, "Materials")
        st.download_button("📊 Download Materials (Excel)", xlsx_buf, file_name=f"materials_{datetime.now().strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    st.subheader("Projects DB")
    # turn projects dict into dataframe for export
    proj_df = pd.DataFrame([(k, json.dumps(v)) for k,v in st.session_state["projects"].items()], columns=["Project Name","Details(JSON)"])
    c3, c4 = st.columns(2)
    with c3:
        word_buf2 = export_to_word_bytes(proj_df, "Projects Database")
        st.download_button("📄 Download Projects (Word)", word_buf2, file_name=f"projects_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with c4:
        xlsx_buf2 = export_to_excel_bytes(proj_df, "Projects")
        st.download_button("📊 Download Projects (Excel)", xlsx_buf2, file_name=f"projects_{datetime.now().strftime('%Y%m%d')}.xlsx", mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

    if st.button("Save current materials to materials_db.json"):
        save_materials_db(st.session_state["materials_df"])
    if st.button("Save projects DB to projects_db.json"):
        save_projects_db(st.session_state["projects"])

# Footer
st.markdown("---")
st.caption(f"Comprehensive estimator — edit prices & consumption to match local BOQs. Files: {MATERIALS_FILE}, {PROJECTS_FILE}. {datetime.utcnow().date()}")

